import os
import qrcode
from skimage import io
from bs4 import BeautifulSoup
import requests
import datetime
from docx import Document
from docx.shared import Inches

TXT_NAME = 'info.txt'  # 存放链接及名称信息的txt
QRCODES_FILENAME = 'qrcode'  # 存放二维码的文件夹

CURRENT_DIR = os.path.abspath(os.path.dirname(__file__))
TXT_PATH = os.path.join(CURRENT_DIR, TXT_NAME)
QRCODE_IMAGE_PATH = os.path.join(CURRENT_DIR, QRCODES_FILENAME)
JOBBJUT_URL = 'https://jobbjut.jysd.com/admin/Campus/Create?source=0&target=navTab&_='

# 从url获取标题


def getTitle(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")

    title = soup.h1.string.strip()
    res = ''.join(filter(str.isalnum, title))  # 去除特殊字符
    return res

# 从txt读取链接信息


def readTxt(doc):
    f = open(TXT_PATH, 'r', encoding='utf-8')
    line = f.readline()
    line = line[:-1]
    while line:
        # read and transfer to qrcode image
        arr = line.split(' ')
        if len(arr) == 1:
            url = line
            title = getTitle(url)
        else:
            url = arr[0]
            title = arr[1]

        url2qrcode(url, title, doc)
        # ---
        line = f.readline()
        line = line[:-1]
    now = datetime.datetime.now().strftime('%Y-%m-%d')
    doc.save('qrcode' + now + '.docx')
    f.close()

# 链接转二维码


def url2qrcode(url, name, doc):
    qr = qrcode.QRCode(
        version=2,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=4,
        border=1
    )  # setting the size 150

    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image()

    # save image
    if not os.path.exists(QRCODE_IMAGE_PATH):
        os.makedirs(QRCODE_IMAGE_PATH)

    imgPath = QRCODE_IMAGE_PATH + '/' + name + '.png'
    img.save(imgPath)

    # save docx
    save_docx(doc, name, imgPath)

# 上传至就业中心后台 (暂未完成)


def upload():
    now = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    title = 'test333'
    content = 'test444'
    formData = {
        'CampusAll[title]': title,
        'CampusAll[color]': '',
        'ignore_company': 'on',
        'CampusAllMore[content]': content,
        'CampusAll[validtime]': '',
        'CampusAll[dateline]': now,
        'CompanyAll[company_name]': '',
        'CompanyAll_id': '',
        'CompanyAll[email]': '',
        'CompanyAll[city]': '',
        'CompanyAll[d_industry2]': '',
        'CompanyAll[d_category]': '',
        'CompanyAll[d_nature]': '',
        'CompanyAll[d_scale]': '',
        'CompanyAllMore[tag]': '',
        'CompanyAllMore[linkman]': '',
        'CompanyAllMore[telephone]': '',
        'CompanyAllMore[contact]': '',
        'CompanyAllMore[description]': '',
        'sendMail_field': '',
        'remark_field': '',
        'vertify': '3',
        'sendMessage_field': '',
        'restrictionsapply': '',
        'ajax': '1'
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36'
    }
    res = requests.post(JOBBJUT_URL, headers=headers, data=formData)
    print(res)


def save_docx(doc, name, img):
    doc.add_paragraph(name)
    doc.add_picture(img, width=Inches(2))


def main():
    doc = Document()
    readTxt(doc)


if __name__ == '__main__':
    main()
